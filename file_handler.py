"""
File handling module for resume-job description comparison.
Handles file uploads and text extraction from various file formats.
"""

import logging
import os
from typing import Dict, Optional, Tuple, Union
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileHandler:
    """
    A class for handling file uploads and text extraction from various formats.
    Supports PDF, DOCX, and plain text files.
    """
    
    def __init__(self):
        """Initialize the FileHandler."""
        self.supported_formats = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.txt': 'text/plain'
        }
        
        self.max_file_size = 10 * 1024 * 1024  # 10MB
    
    def validate_file(self, file_path: str, file_type: Optional[str] = None) -> Tuple[bool, str]:
        """
        Validate if a file can be processed.
        
        Args:
            file_path (str): Path to the file
            file_type (str, optional): Expected file type
            
        Returns:
            Tuple[bool, str]: (is_valid, error_message)
        """
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                return False, f"File not found: {file_path}"
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                return False, f"File too large: {file_size / (1024*1024):.1f}MB (max: {self.max_file_size / (1024*1024)}MB)"
            
            # Check file extension
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in self.supported_formats:
                return False, f"Unsupported file format: {file_ext}. Supported: {', '.join(self.supported_formats.keys())}"
            
            # Check file type if specified
            if file_type and file_ext != file_type.lower():
                return False, f"File type mismatch. Expected: {file_type}, got: {file_ext}"
            
            return True, ""
            
        except Exception as e:
            logger.error(f"Error validating file {file_path}: {str(e)}")
            return False, f"Validation error: {str(e)}"
    
    def extract_text_from_pdf(self, file_path: str) -> Tuple[bool, Union[str, str]]:
        """
        Extract text from a PDF file.
        
        Args:
            file_path (str): Path to the PDF file
            
        Returns:
            Tuple[bool, Union[str, str]]: (success, text_or_error_message)
        """
        try:
            import pdfplumber
            
            # Validate file
            is_valid, error_msg = self.validate_file(file_path, '.pdf')
            if not is_valid:
                return False, error_msg
            
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        logger.info(f"Processed page {page_num + 1}")
                    except Exception as e:
                        logger.warning(f"Error processing page {page_num + 1}: {str(e)}")
                        continue
            
            if not text.strip():
                return False, "No text could be extracted from the PDF"
            
            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return True, text.strip()
            
        except ImportError:
            return False, "pdfplumber library not installed. Please install it using: pip install pdfplumber"
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            return False, f"Error extracting text: {str(e)}"
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[bool, Union[str, str]]:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path (str): Path to the DOCX file
            
        Returns:
            Tuple[bool, Union[str, str]]: (success, text_or_error_message)
        """
        try:
            from docx import Document
            
            # Validate file
            is_valid, error_msg = self.validate_file(file_path, '.docx')
            if not is_valid:
                return False, error_msg
            
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            if not text.strip():
                return False, "No text could be extracted from the DOCX file"
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return True, text.strip()
            
        except ImportError:
            return False, "python-docx library not installed. Please install it using: pip install python-docx"
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return False, f"Error extracting text: {str(e)}"
    
    def extract_text_from_txt(self, file_path: str) -> Tuple[bool, Union[str, str]]:
        """
        Extract text from a plain text file.
        
        Args:
            file_path (str): Path to the text file
            
        Returns:
            Tuple[bool, Union[str, str]]: (success, text_or_error_message)
        """
        try:
            # Validate file
            is_valid, error_msg = self.validate_file(file_path, '.txt')
            if not is_valid:
                return False, error_msg
            
            # Read text file with proper encoding detection
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
                except Exception as e:
                    logger.warning(f"Error reading with encoding {encoding}: {str(e)}")
                    continue
            
            if not text.strip():
                return False, "No text could be extracted from the file"
            
            logger.info(f"Successfully extracted {len(text)} characters from text file")
            return True, text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from text file {file_path}: {str(e)}")
            return False, f"Error extracting text: {str(e)}"
    
    def extract_text(self, file_path: str) -> Tuple[bool, Union[str, str]]:
        """
        Extract text from a file based on its extension.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            Tuple[bool, Union[str, str]]: (success, text_or_error_message)
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            return False, f"Unsupported file format: {file_ext}"
    
    def get_file_info(self, file_path: str) -> Dict[str, any]:
        """
        Get information about a file.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            Dict[str, any]: File information
        """
        try:
            if not os.path.exists(file_path):
                return {'error': 'File not found'}
            
            file_stat = os.stat(file_path)
            file_ext = Path(file_path).suffix.lower()
            
            return {
                'name': os.path.basename(file_path),
                'path': file_path,
                'extension': file_ext,
                'size_bytes': file_stat.st_size,
                'size_mb': round(file_stat.st_size / (1024 * 1024), 2),
                'created': file_stat.st_ctime,
                'modified': file_stat.st_mtime,
                'is_readable': os.access(file_path, os.R_OK),
                'supported_format': file_ext in self.supported_formats
            }
            
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {str(e)}")
            return {'error': f'Error getting file info: {str(e)}'}
    
    def save_text_to_file(self, text: str, output_path: str) -> Tuple[bool, str]:
        """
        Save extracted text to a file.
        
        Args:
            text (str): Text to save
            output_path (str): Output file path
            
        Returns:
            Tuple[bool, str]: (success, error_message)
        """
        try:
            # Create directory if it doesn't exist
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            # Save text
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(text)
            
            logger.info(f"Text saved to {output_path}")
            return True, ""
            
        except Exception as e:
            logger.error(f"Error saving text to {output_path}: {str(e)}")
            return False, f"Error saving file: {str(e)}"
    
    def batch_process_files(self, file_paths: list, output_dir: str = None) -> Dict[str, any]:
        """
        Process multiple files and extract text from each.
        
        Args:
            file_paths (list): List of file paths to process
            output_dir (str, optional): Directory to save extracted text files
            
        Returns:
            Dict[str, any]: Results of batch processing
        """
        results = {
            'total_files': len(file_paths),
            'successful': 0,
            'failed': 0,
            'results': {}
        }
        
        for file_path in file_paths:
            try:
                # Extract text
                success, text_or_error = self.extract_text(file_path)
                
                if success:
                    results['successful'] += 1
                    results['results'][file_path] = {
                        'status': 'success',
                        'text_length': len(text_or_error),
                        'text_preview': text_or_error[:200] + "..." if len(text_or_error) > 200 else text_or_error
                    }
                    
                    # Save to output directory if specified
                    if output_dir:
                        output_filename = Path(file_path).stem + "_extracted.txt"
                        output_path = os.path.join(output_dir, output_filename)
                        save_success, save_error = self.save_text_to_file(text_or_error, output_path)
                        
                        if save_success:
                            results['results'][file_path]['saved_to'] = output_path
                        else:
                            results['results'][file_path]['save_error'] = save_error
                    
                else:
                    results['failed'] += 1
                    results['results'][file_path] = {
                        'status': 'failed',
                        'error': text_or_error
                    }
                    
            except Exception as e:
                results['failed'] += 1
                results['results'][file_path] = {
                    'status': 'error',
                    'error': str(e)
                }
                logger.error(f"Error processing file {file_path}: {str(e)}")
        
        return results
    
    def cleanup_temp_files(self, temp_files: list) -> None:
        """
        Clean up temporary files.
        
        Args:
            temp_files (list): List of temporary file paths to remove
        """
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
                    logger.info(f"Removed temporary file: {temp_file}")
            except Exception as e:
                logger.warning(f"Could not remove temporary file {temp_file}: {str(e)}")
    
    def get_supported_formats_info(self) -> Dict[str, any]:
        """
        Get information about supported file formats.
        
        Returns:
            Dict[str, any]: Information about supported formats
        """
        return {
            'formats': self.supported_formats,
            'max_file_size_mb': self.max_file_size / (1024 * 1024),
            'extensions': [ext.replace('.', '') for ext in self.supported_formats.keys()]
        }
